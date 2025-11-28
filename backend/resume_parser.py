import os
import PyPDF2
import docx
import json
from typing import Dict, Any
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from config import Config

class ResumeParser:
    def __init__(self):
        self.llm = ChatOpenAI(
            model="gpt-3.5-turbo",
            temperature=0,
            openai_api_key=Config.OPENAI_API_KEY
        )
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF has pages
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF file has no pages")
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                # Check if any text was extracted
                if not text.strip():
                    raise ValueError("PDF file contains no extractable text (might be image-based)")
                    
        except PyPDF2.errors.PdfReadError as e:
            raise ValueError(f"Invalid or corrupted PDF file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
        
        return text.strip()
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            
            # Check if document has content
            if len(doc.paragraphs) == 0:
                raise ValueError("DOCX file has no paragraphs")
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            # Check if any text was extracted
            if not text.strip():
                raise ValueError("DOCX file contains no extractable text")
                
        except docx.opc.exceptions.PackageNotFoundError:
            raise ValueError("Invalid DOCX file format")
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
        
        return text.strip()
    
    def extract_text(self, file_path: str) -> str:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def parse_resume_with_llm(self, text: str) -> Dict[str, Any]:
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert resume parser. Extract the following information from the resume text and return it as a JSON object.
            
Required fields:
- name: Full name of the candidate
- email: Email address
- phone: Phone number
- company: Current or most recent company
- designation: Current or most recent job title/designation
- skills: Array of skills (technical and soft skills)

For each extracted field, also provide a confidence score (0.0 to 1.0) indicating how confident you are about the extraction.

Return the response in this exact JSON format:
{{
    "data": {{
        "name": "extracted name or null",
        "email": "extracted email or null",
        "phone": "extracted phone or null",
        "company": "extracted company or null",
        "designation": "extracted designation or null",
        "skills": ["skill1", "skill2", ...]
    }},
    "confidence_scores": {{
        "name": 0.0-1.0,
        "email": 0.0-1.0,
        "phone": 0.0-1.0,
        "company": 0.0-1.0,
        "designation": 0.0-1.0,
        "skills": 0.0-1.0
    }}
}}

If a field cannot be found, set it to null and give it a confidence score of 0.0.
"""),
            ("user", "Resume text:\n\n{resume_text}")
        ])
        
        try:
            chain = prompt | self.llm
            response = chain.invoke({"resume_text": text})
            
            # Parse JSON response
            try:
                result = json.loads(response.content)
            except json.JSONDecodeError as e:
                raise ValueError(f"LLM returned invalid JSON: {str(e)}")
            
            # Validate structure
            if not isinstance(result, dict):
                raise ValueError("LLM response is not a JSON object")
            
            if 'data' not in result:
                raise ValueError("LLM response missing 'data' field")
            
            data = result.get('data', {})
            
            # Ensure at least one identifying field is present
            has_name = data.get('name') and str(data.get('name')).strip() and str(data.get('name')).lower() != 'null'
            has_email = data.get('email') and str(data.get('email')).strip() and str(data.get('email')).lower() != 'null'
            has_phone = data.get('phone') and str(data.get('phone')).strip() and str(data.get('phone')).lower() != 'null'
            
            if not (has_name or has_email or has_phone):
                raise ValueError("Could not extract any identifying information (name, email, or phone) from resume. The resume may be corrupted, image-based, or have invalid content.")
            
            # Ensure confidence_scores exists
            if 'confidence_scores' not in result:
                result['confidence_scores'] = {
                    "name": 0.5,
                    "email": 0.5,
                    "phone": 0.5,
                    "company": 0.5,
                    "designation": 0.5,
                    "skills": 0.5
                }
            
            # Ensure skills is a list
            if not isinstance(data.get('skills'), list):
                data['skills'] = []
            
            return result
            
        except ValueError as e:
            # Re-raise validation errors
            raise
        except Exception as e:
            # Return default structure for parsing failures
            raise Exception(f"Error parsing resume with LLM: {str(e)}")
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Main method to parse resume and extract information"""
        # Validate file exists
        if not os.path.exists(file_path):
            raise ValueError(f"Resume file not found: {file_path}")
        
        # Validate file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError("Resume file is empty")
        
        # Extract text
        text = self.extract_text(file_path)
        
        # Validate extracted text has minimum length
        if len(text) < 50:
            raise ValueError("Resume text is too short (less than 50 characters). Please ensure the file contains valid resume content.")
        
        # Parse with LLM
        result = self.parse_resume_with_llm(text)
        
        return result
